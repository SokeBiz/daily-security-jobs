"""
DOCX report generator for cybersecurity jobs.

Produces a formatted .docx with job summaries, locations, descriptions,
and apply links.
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import truncate_text


def generate_docx(
    matched_jobs: list[dict[str, Any]],
    lookback_hours: int,
    output_dir: str | Path = "/tmp",
    platform_counts: dict[str, int] | None = None,
) -> Path:
    """Generate a .docx report with matched cybersecurity jobs.

    Args:
        matched_jobs: List of normalized job dicts that passed all filters.
        lookback_hours: Lookback window for the summary line.
        output_dir: Directory to save the .docx file.
        platform_counts: Dict of platform_name -> company_count for footer.

    Returns:
        Path to the generated .docx file.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_path = output_dir / f"security-jobs-{timestamp}.docx"

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Title ──
    title = doc.add_heading("Security Jobs Digest", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Summary ──
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(
        f"Jobs posted in the last {lookback_hours} hours | "
        f"{len(matched_jobs)} positions found | "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    if not matched_jobs:
        doc.add_paragraph("No matching security jobs found in this window.")
        doc.save(str(output_path))
        return output_path

    # ── Jobs ──
    for i, job in enumerate(matched_jobs, 1):
        doc.add_heading(f"{i}. {job['title']}", level=2)

        # Company + Source tag
        info = doc.add_paragraph()
        run = info.add_run(f"\U0001f3e2 {job['company'].title()}")
        run.bold = True
        run.font.size = Pt(10)

        run = info.add_run(f"  |  \U0001f4cb {job['source']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 180)

        # Details
        details = [
            ("\U0001f4cd Location", job["location"]),
            ("\U0001f3f7\ufe0f Department", job["department"] or "N/A"),
            (
                "\U0001f550 Published",
                job["published_at"].strftime("%Y-%m-%d %H:%M UTC")
                if job["published_at"]
                else "Unknown",
            ),
            ("\U0001f3e0 Workplace", job.get("workplace_type", "N/A") or "N/A"),
        ]

        for label, value in details:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(value)
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

        # Description snippet
        if job.get("description"):
            desc_text = truncate_text(job["description"], 400)
            p = doc.add_paragraph()
            run = p.add_run("\U0001f4dd Description: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(desc_text)
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)

        # Apply link
        if job.get("apply_url"):
            p = doc.add_paragraph()
            run = p.add_run("\U0001f517 Apply: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(job["apply_url"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 76, 153)

        # Separator
        doc.add_paragraph("\u2500" * 60)

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Job sources: Ashby, Greenhouse, Lever, Personio, Recruitee, Amazon, Workday, Accenture"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    if platform_counts:
        counts_str = " \u00b7 ".join(
            f"{name}: {count}" for name, count in platform_counts.items()
        )
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"Companies scanned: {counts_str}")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(str(output_path))
    return output_path
